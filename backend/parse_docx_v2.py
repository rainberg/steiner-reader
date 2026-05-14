#!/usr/bin/env python3
"""
改进版 docx 解析器：
1. 更全面的章节标题识别（序数词 + 日期格式 + 编号格式）
2. 过滤非正文内容（Inhalt/Vorrede/Einleitung 等）
3. 更准确的段落和句子拆分
4. 输出 JSON 到 stdout
"""
import sys, json, re
from docx import Document

def main():
    path = sys.argv[1]
    doc = Document(path)
    
    # ===== 标题识别正则 =====
    # 德语序数词讲座标题：ERSTER VORTRAG, ZWEITER VORTRAG 等
    ORDINAL_RE = re.compile(
        r"(ERSTER|ZWEITER|DRITTER|VIERTER|FUENFTER|FÜNFTER|SECHSTER|SIEBENTER|SIEBTER|"
        r"ACHTER|NEUNTER|ZEHNTER|ELFTER|ZWOELFTER|ZWÖLFTER|DREIZEHNTER|VIERZEHNTER|"
        r"FUENFZEHNTER|FÜNFZEHNTER|SECHZEHNTER|SIEBZEHNTER|ACHTZEHNTER|NEUNZEHNTER|"
        r"ZWANZIGSTER|ERSTE|ZWEITE|DRITTE|VIERTE|FÜNFTE|SECHSTE|SIEBENTE|"
        r"ERSTES|ZWEITES|DRITTES|VIERTES|FÜNFTES)\s+(VORTRAG|KAPITEL|TEIL|HAUPTSTÜCK|BAND|VORLESUNG)",
        re.IGNORECASE
    )
    
    # 带编号的讲座：1. Vortrag, Vortrag 1, Kapitel 1 等
    NUMBERED_RE = re.compile(
        r"^(\d{1,2})\.\s*(Vortrag|VORTRAG|Kapitel|KAPITEL|Vorlesung|VORLESUNG|"
        r"Teil|TEIL|Übung|ÜBUNG|Predigt|PREDIGT)",
        re.IGNORECASE
    )
    
    # 日期格式标题：Stuttgart, 16. Juni 1921 或 Dornach, 13. Oktober 1920
    DATE_TITLE_RE = re.compile(
        r"^([A-ZÄÖÜa-zäöüß][A-ZÄÖÜa-zäöüß\s\-/]+?),\s+"
        r"(\d{1,2}\.\s*(?:Januar|Februar|März|Marz|April|Mai|Juni|Juli|"
        r"August|September|Oktober|November|Dezember)\s+\d{4})"
    )
    
    # 罗马数字章节：I., II., III., IV. 等（单独一行，后面跟内容描述）
    ROMAN_RE = re.compile(
        r"^(I{1,3}|IV|V|VI{1,3}|IX|X{1,3}|XI|XII|XIII|XIV|XV|XVI{1,3}|XIX|XX)"
        r"[\.\s]+([A-ZÄÖÜ][a-zäöüß]+)",
        re.IGNORECASE
    )
    
    # ===== 非正文章节识别 =====
    SKIP_PATTERNS = [
        re.compile(r"^(INHALT|Inhalt|CONTENTS|TABLE OF CONTENTS|INHALTSVERZEICHNIS)\s*$", re.IGNORECASE),
        re.compile(r"^(VORREDE|Vorrede|PREFACE|FOREWORD)\s*$", re.IGNORECASE),
        re.compile(r"^(EINLEITUNG|Einleitung|INTRODUCTION)\s*$", re.IGNORECASE),
        re.compile(r"^(VORWORT|Vorwort)\s*$", re.IGNORECASE),
        re.compile(r"^(VORBEMERKUNG|Vorbemerkung|PRELIMINARY REMARKS)\s*$", re.IGNORECASE),
        re.compile(r"^(REGISTER|Register|INDEX|Index|LITERATURVERZEICHNIS)\s*$", re.IGNORECASE),
        re.compile(r"^(ANHANG|Anhang|APPENDIX)\s*$", re.IGNORECASE),
        re.compile(r"^(NACHWORT|Nachwort|AFTERWORD)\s*$", re.IGNORECASE),
    ]
    
    # 目录行特征：标题后跟页码数字
    TOC_LINE_RE = re.compile(r".+\.{2,}\s*\d+\s*$|.+[\s\.]{3,}\d+\s*$")
    
    def is_toc_line(text):
        """检测是否是目录行（标题...页码）"""
        return bool(TOC_LINE_RE.match(text.strip()))
    
    def clean_text(t):
        """清理文本：合并断字、去除脚注标记、规范空白"""
        t = re.sub(r"[¹²³⁴⁵⁶⁷⁸⁹⁰]+|\(\d+\)", "", t)
        t = re.sub(r"-\n(\w)", r"\1", t)  # 合并断字
        t = re.sub(r"\s+", " ", t)
        return t.strip()
    
    def split_paragraphs(text):
        """拆分段落：按双换行或首行缩进"""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # 按空行拆分
        parts = re.split(r"\n\s*\n", text)
        paragraphs = []
        for p in parts:
            p = p.strip()
            if len(p) > 20:  # 过滤太短的段落（可能是页码等）
                paragraphs.append(p)
        return paragraphs
    
    def split_sentences(text):
        """拆分句子"""
        text = text.strip()
        if len(text) < 10:
            return []
        # 按句号、问号、感叹号后跟大写字母拆分
        parts = re.split(r'(?<=[.!?])\s+(?=[A-ZÄÖÜ"\'\(\])', text)
        sentences = [s.strip() for s in parts if s.strip() and len(s.strip()) > 5]
        return sentences if sentences else [text]
    
    # ===== 解析文档 =====
    # 获取书名
    title_de = ""
    ga_number = ""
    for para in doc.paragraphs[:20]:
        t = para.text.strip()
        if t and len(t) > 3:
            title_de = t
            m = re.search(r"GA\s*(\d+)", t, re.IGNORECASE)
            if m:
                ga_number = f"GA {m.group(1)}"
            break
    
    # 从元数据获取
    props = doc.core_properties
    if props.title and not title_de:
        title_de = props.title
    
    # ===== 按段落遍历，识别章节 =====
    lectures = []
    cur_lecture = None
    lorder = 0
    pcounter = 0
    in_skip_section = False
    skip_until_next_header = False
    toc_detected = False
    content_started = False  # 是否已进入正文
    
    # 统计连续的目录行
    toc_line_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        cleaned = clean_text(text)
        if not cleaned:
            continue
        
        # ===== 检测是否是章节标题 =====
        is_header = False
        title = ""
        location = ""
        date = ""
        
        # 检查序数词标题
        m = ORDINAL_RE.search(cleaned)
        if m:
            is_header = True
            title = cleaned[:100]
            # 尝试提取地点和日期
            dm = DATE_TITLE_RE.search(cleaned)
            if dm:
                location = dm.group(1).strip()
                date = dm.group(2).strip()
        
        # 检查编号标题
        if not is_header:
            m = NUMBERED_RE.match(cleaned)
            if m:
                is_header = True
                title = cleaned[:100]
                dm = DATE_TITLE_RE.search(cleaned)
                if dm:
                    location = dm.group(1).strip()
                    date = dm.group(2).strip()
        
        # 检查日期格式标题
        if not is_header:
            m = DATE_TITLE_RE.match(cleaned)
            if m and len(cleaned) < 120:
                is_header = True
                title = cleaned[:100]
                location = m.group(1).strip()
                date = m.group(2).strip()
        
        # 检查罗马数字标题
        if not is_header:
            m = ROMAN_RE.match(cleaned)
            if m and len(cleaned) < 100:
                is_header = True
                title = cleaned[:100]
        
        # ===== 检测非正文区域 =====
        skip_match = False
        for pat in SKIP_PATTERNS:
            if pat.match(cleaned):
                in_skip_section = True
                skip_until_next_header = True
                skip_match = True
                break
        
        if skip_match:
            continue
        
        # 检测目录行
        if is_toc_line(cleaned):
            toc_line_count += 1
            if toc_line_count > 3:
                toc_detected = True
            continue
        else:
            if toc_line_count > 0 and not toc_detected:
                toc_line_count = 0
        
        if toc_detected:
            # 目录结束后，遇到第一个真正的标题才开始
            if is_header:
                toc_detected = False
                content_started = True
            else:
                continue
        
        # ===== 处理章节标题 =====
        if is_header:
            # 如果之前在跳过区域，现在遇到新标题就恢复
            if skip_until_next_header:
                skip_until_next_header = False
                in_skip_section = False
            
            content_started = True
            
            cur_lecture = {
                "title_de": title,
                "location": location,
                "date": date,
                "order_index": lorder,
                "paragraphs": []
            }
            lectures.append(cur_lecture)
            lorder += 1
            pcounter = 0
            continue
        
        # 跳过非正文区域的内容
        if in_skip_section or not content_started:
            continue
        
        # ===== 处理正文内容 =====
        if cur_lecture is None:
            # 还没遇到标题，创建一个默认章节
            cur_lecture = {
                "title_de": title_de or "Gesamtwerk",
                "location": "",
                "date": "",
                "order_index": 0,
                "paragraphs": []
            }
            lectures.append(cur_lecture)
            lorder = 1
            pcounter = 0
        
        # 处理段落
        paras = split_paragraphs(cleaned)
        for pt in paras:
            if re.match(r"^\s*\d{1,3}\s*$", pt):  # 纯页码
                continue
            if is_toc_line(pt):  # 目录行
                continue
            
            sents = split_sentences(pt)
            para_obj = {"order_index": pcounter, "sentences": []}
            pcounter += 1
            
            for idx, st in enumerate(sents):
                if len(st) > 5:
                    para_obj["sentences"].append({"text_de": st, "order_index": idx})
            
            if para_obj["sentences"]:
                cur_lecture["paragraphs"].append(para_obj)
    
    # doc.close()  # python-docx does not need explicit close
    
    result = {
        "title_de": title_de,
        "ga_number": ga_number,
        "lectures": lectures
    }
    
    json.dump(result, sys.stdout, ensure_ascii=False)

if __name__ == "__main__":
    main()
