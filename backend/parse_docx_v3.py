#!/usr/bin/env python3
"""
docx 解析器 v3：用段落样式（Heading 1）识别章节，过滤非正文内容
输出 JSON 到 stdout
"""
import sys, json, re
from docx import Document

def main():
    path = sys.argv[1]
    doc = Document(path)
    
    # ===== 非正文章节标题 =====
    SKIP_TITLES = {
        "inhalt", "contents", "table of contents", "inhaltverzeichnis",
        "vorrede", "preface", "foreword",
        "vorwort", "vorwort zur",
        "einleitung", "introduction",
        "vorbemerkung", "preliminary remarks",
        "register", "index", "literaturverzeichnis",
        "anhang", "appendix",
        "nachwort", "afterword",
    }
    
    def is_skip_title(title):
        t = title.lower().strip()
        for skip in SKIP_TITLES:
            if t.startswith(skip):
                return True
        return False
    
    def clean_text(t):
        t = re.sub(r"[¹²³⁴⁵⁶⁷⁸⁹⁰]+|\(\d+\)", "", t)
        t = re.sub(r"-\n(\w)", r"\1", t)
        t = re.sub(r"\s+", " ", t)
        return t.strip()
    
    def split_sentences(text):
        text = text.strip()
        if len(text) < 10:
            return []
        parts = re.split(r'(?<=[.!?])\s+(?=[A-ZÄÖÜ"\'])', text)
        sentences = [s.strip() for s in parts if s.strip() and len(s.strip()) > 5]
        return sentences if sentences else [text]
    
    # ===== 获取书名 =====
    title_de = ""
    ga_number = ""
    for para in doc.paragraphs[:15]:
        t = para.text.strip()
        if t and len(t) > 5 and not t.startswith("http"):
            title_de = t
            m = re.search(r"GA\s*(\d+)", t, re.IGNORECASE)
            if m:
                ga_number = "GA %s" % m.group(1)
            break
    
    props = doc.core_properties
    if props.title and props.title.strip():
        title_de = props.title.strip()
    
    # ===== 遍历段落，用样式识别章节 =====
    lectures = []
    cur_lecture = None
    lorder = 0
    pcounter = 0
    skip_current_section = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name or "heading" in style_name
        
        # ===== 处理章节标题 =====
        if is_heading:
            cleaned = clean_text(text)
            if not cleaned:
                continue
            
            # 检查是否需要跳过
            if is_skip_title(cleaned):
                skip_current_section = True
                cur_lecture = None
                continue
            
            skip_current_section = False
            
            # 提取地点和日期
            location = ""
            date = ""
            dm = re.search(
                r"([A-ZÄÖÜa-zäöüß][A-ZÄÖÜa-zäöüß\s\-/]+?),\s+"
                r"(\d{1,2}\.\s*(?:Januar|Februar|März|April|Mai|Juni|Juli|"
                r"August|September|Oktober|November|Dezember)\s+\d{4})",
                cleaned
            )
            if dm:
                location = dm.group(1).strip()
                date = dm.group(2).strip()
            
            cur_lecture = {
                "title_de": cleaned[:200],
                "location": location,
                "date": date,
                "order_index": lorder,
                "paragraphs": []
            }
            lectures.append(cur_lecture)
            lorder += 1
            pcounter = 0
            continue
        
        # 跳过当前非正文区域
        if skip_current_section:
            continue
        
        # ===== 处理正文 =====
        cleaned = clean_text(text)
        if not cleaned or len(cleaned) < 15:
            continue
        
        # 跳过页码
        if re.match(r"^\s*\d{1,4}\s*$", cleaned):
            continue
        
        # 跳过目录行（标题...页码）
        if re.match(r".+\.{2,}\s*\d+\s*$", cleaned):
            continue
        
        # 跳过元数据行
        if cleaned.startswith("Rudolf Steiner Online") or cleaned.startswith("Seitennummern"):
            continue
        if re.match(r"^(Steiner, Rudolf|Rudolf Steiner)\s*:?\s*$", cleaned):
            continue
        if re.match(r"^\w+ \(CH\)\.", cleaned):
            continue
        
        # 如果还没遇到标题，创建默认章节
        if cur_lecture is None:
            cur_lecture = {
                "title_de": title_de or "Gesamtwerk",
                "location": "",
                "date": "",
                "order_index": 0,
                "paragraphs": []
            }
            lectures.append(cur_lecture)
            lorder = 1
            pcounter = 0
        
        # 拆分句子
        sents = split_sentences(cleaned)
        if not sents:
            continue
        
        para_obj = {"order_index": pcounter, "sentences": []}
        pcounter += 1
        
        for idx, st in enumerate(sents):
            if len(st) > 5:
                para_obj["sentences"].append({"text_de": st, "order_index": idx})
        
        if para_obj["sentences"]:
            cur_lecture["paragraphs"].append(para_obj)
    
    result = {
        "title_de": title_de,
        "ga_number": ga_number,
        "lectures": lectures
    }
    
    json.dump(result, sys.stdout, ensure_ascii=False)

if __name__ == "__main__":
    main()
