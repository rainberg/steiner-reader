#!/usr/bin/env python3
"""
docx 解析器 v5 - 混合模式：
1. 优先用 Heading 1 样式识别章节
2. 若无 Heading 1，回退到正则模式（erster vortrag, Kapitel I. 等）
3. 智能过滤目录区域（页码行）和非正文章节
"""
import sys, json, re, os
from docx import Document

def main():
    path = sys.argv[1]
    doc = Document(path)
    
    basename = os.path.basename(path)
    ga_match = re.search(r"GA(\d+)", basename, re.IGNORECASE)
    ga_number = "GA %s" % ga_match.group(1) if ga_match else ""
    
    # ===== 正则模式 =====
    # 德语序数词讲座
    LECTURE_ORDINAL = re.compile(
        r"^((?:erster|zweiter|dritter|vierter|fünfter|sechster|siebenter|"
        r"achter|neunter|zehnter|elfter|zwölfter|dreizehnter|vierzehnter|"
        r"fünfzehnter|sechzehnter|siebzehnter|achtzehnter|neunzehnter|"
        r"zwanzigster)\s+vortrag)",
        re.IGNORECASE
    )
    
    # 德语大写序数词讲座：ERSTER VORTRAG
    LECTURE_UPPER = re.compile(
        r"^((?:ERSTER|ZWEITER|DRITTER|VIERTER|FÜNFTER|SECHSTER|SIEBENTER|"
        r"ACHTER|NEUNTER|ZEHNTER|ELFTER|ZWÖLFTER|DREIZEHNTER|VIERZEHNTER|"
        r"FÜNFZEHNTER|SECHZEHNTER|SIEBZEHNTER|ACHTZEHNTER|NEUNZEHNTER|"
        r"ZWANZIGSTER)\s+VORTRAG)",
        re.IGNORECASE
    )
    
    # 编号讲座：1. Vortrag
    LECTURE_NUM = re.compile(r"^(\d{1,2})\.\s*Vortrag\b", re.IGNORECASE)
    
    # Kapitel
    CHAPTER_RE = re.compile(r"^Kapitel\s+(I{1,4}|V?I{0,4}|IX|X{0,3}|XI{0,4})\.?\s*$", re.IGNORECASE)
    
    # 跳过的标题
    SKIP_TITLES = [
        "inhalt", "contents", "table of contents", "vorrede", "preface",
        "vorwort", "einleitung", "introduction", "vorbemerkung",
        "register", "index", "anhang", "appendix", "nachwort", "nachwort zum",
    ]
    
    def is_skip_title(title):
        t = title.lower().strip()
        for prefix in SKIP_TITLES:
            if t.startswith(prefix):
                return True
        return False
    
    # 目录行检测：标题 + 多个空格/tab + 数字
    TOC_LINE = re.compile(r"^.+?[\t ]{2,}\d{1,4}\s*$")
    
    def clean_text(t):
        t = re.sub(r"[¹²³⁴⁵⁶⁷⁸⁹⁰]+|\(\d+\)", "", t)
        t = re.sub(r"-\n(\w)", r"\1", t)
        t = re.sub(r"\s+", " ", t)
        return t.strip()
    
    def split_sentences(text):
        text = text.strip()
        if len(text) < 10:
            return []
        parts = re.split(r'(?<=[.!?])\s+(?=[A-ZÄÖÜ"\'])', text)
        sentences = [s.strip() for s in parts if s.strip() and len(s.strip()) > 5]
        return sentences if sentences else [text]
    
    # ===== 检测是否有 Heading 1 =====
    has_h1 = False
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1":
            has_h1 = True
            break
    
    # ===== 遍历段落 =====
    lectures = []
    cur_lecture = None
    lorder = 0
    pcounter = 0
    skip_section = False
    title_de = ""
    in_toc = True  # 开始时假设在目录区域
    toc_line_count = 0
    long_para_seen = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        cleaned = clean_text(text)
        if not cleaned:
            continue
        
        style_name = para.style.name if para.style else ""
        
        # ===== 检测是否离开目录区域 =====
        if in_toc:
            # 遇到长段落 (>150字符) 说明进入正文
            if len(cleaned) > 150:
                in_toc = False
                long_para_seen = True
            # 目录行跳过
            elif TOC_LINE.match(cleaned):
                toc_line_count += 1
                continue
            elif toc_line_count > 3:
                # 目录结束后遇到非目录行，但也不够长
                # 等待第一个章节标题
                pass
        
        if in_toc:
            continue
        
        # ===== Heading 1 模式 =====
        is_h1 = has_h1 and style_name == "Heading 1"
        
        # ===== 正则模式 =====
        is_regex_header = False
        header_title = ""
        
        if not has_h1 or not is_h1:
            # 检查序数词讲座
            m = LECTURE_ORDINAL.match(cleaned) or LECTURE_UPPER.match(cleaned)
            if m:
                is_regex_header = True
                header_title = cleaned[:100]
            
            # 检查编号讲座
            if not is_regex_header:
                m = LECTURE_NUM.match(cleaned)
                if m:
                    is_regex_header = True
                    header_title = cleaned[:100]
            
            # 检查 Kapitel
            if not is_regex_header:
                m = CHAPTER_RE.match(cleaned)
                if m:
                    is_regex_header = True
                    header_title = cleaned[:100]
        
        is_header = is_h1 or is_regex_header
        
        # ===== 处理标题 =====
        if is_header:
            title_text = cleaned if is_h1 else header_title
            
            # 跳过非正文章节
            if is_skip_title(title_text):
                skip_section = True
                cur_lecture = None
                continue
            
            skip_section = False
            
            if not title_de:
                title_de = title_text
            
            # 提取地点和日期
            location = ""
            date = ""
            dm = re.search(
                r"([A-ZÄÖÜa-zäöüß][A-ZÄÖÜa-zäöüß\s\-/]+?),\s+"
                r"(\d{1,2}\.\s*(?:Januar|Februar|März|April|Mai|Juni|Juli|"
                r"August|September|Oktober|November|Dezember)\s+\d{4})",
                title_text
            )
            if dm:
                location = dm.group(1).strip()
                date = dm.group(2).strip()
            
            cur_lecture = {
                "title_de": title_text[:200],
                "location": location,
                "date": date,
                "order_index": lorder,
                "paragraphs": []
            }
            lectures.append(cur_lecture)
            lorder += 1
            pcounter = 0
            continue
        
        if skip_section:
            continue
        
        # ===== 处理正文 =====
        # 跳过太短的段落
        if len(cleaned) < 15:
            continue
        
        # 跳过纯页码
        if re.match(r"^\s*\d{1,4}\s*$", cleaned):
            continue
        
        # 跳过目录行
        if TOC_LINE.match(cleaned):
            continue
        
        # 跳过元数据行
        skip_meta = False
        for prefix in ["Rudolf Steiner Online", "Seitennummern", "Anmerkungen, angegeben",
                       "Nach vom Vortragenden", "Zu den Veröffentlichungen",
                       "Die Grundlage der anthroposophisch", "Nach dem Tode von Marie Steiner",
                       "Wir übergeben der Öffentlichkeit"]:
            if cleaned.startswith(prefix):
                skip_meta = True
                break
        if re.match(r"^(Steiner, Rudolf:|Rudolf Steiner|Fragenbeantwort)\s*$", cleaned, re.IGNORECASE):
            skip_meta = True
        if skip_meta:
            continue
        
        # 还没遇到标题，创建默认章节
        if cur_lecture is None:
            cur_lecture = {
                "title_de": title_de or ga_number or "Gesamtwerk",
                "location": "",
                "date": "",
                "order_index": 0,
                "paragraphs": []
            }
            lectures.append(cur_lecture)
            lorder = 1
            pcounter = 0
        
        # 拆分句子
        sents = split_sentences(cleaned)
        if not sents:
            continue
        
        para_obj = {"order_index": pcounter, "sentences": []}
        pcounter += 1
        
        for idx, st in enumerate(sents):
            if len(st) > 5:
                para_obj["sentences"].append({"text_de": st, "order_index": idx})
        
        if para_obj["sentences"]:
            cur_lecture["paragraphs"].append(para_obj)
    
    if not title_de:
        title_de = ga_number or "Unbekannt"
    
    result = {
        "title_de": title_de,
        "ga_number": ga_number,
        "lectures": lectures
    }
    
    json.dump(result, sys.stdout, ensure_ascii=False)

if __name__ == "__main__":
    main()
